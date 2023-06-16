## Opening of Python Packages to run the program ###
import streamlit as st
import numpy as np
import os
import tempfile
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Inches
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.table import _Cell
from docx.oxml.ns import qn

from datetime import date
import string


##########################################
#----------------------------------------#
#              Title Panel               #
#----------------------------------------#
##########################################

st.set_page_config(layout='wide')
st.header("Master Batch Record Drafter")
st.caption("[Not Ready]")

##########################################
#----------------------------------------#
#             ReadMe Panel               #
#----------------------------------------#
##########################################

st.markdown("### ***README***")
with st.expander("User Guide",expanded=True):
#   st.caption("Step 0: Check the box to prevent auto reset")
#   st.caption("Step 1: Type Document Name in Sidebar ")
#   st.caption("Step 2: Type Output Name in Sidebar ")
#   st.caption("Step 3: Type Batch Number in Sidebar ")
#   st.caption("Step 4: Type Your Name in Sidebar ")
#   st.caption("Step 5: Select Packaging Process(s) - [List of Processes] ")
#   st.caption("Step 6: Select Specific Steps - [Process Control Panel] ")
   st.caption("Coming soon...")

if 'writing draft' not in st.session_state:
   st.session_state['writing draft']=False

#---------- Sidebar Setup
stateholder = st.sidebar.checkbox("Step 0: Check this box to prevent unwanted rerun")
if stateholder:
   st.session_state['writing draft']=True

customerName = st.sidebar.text_input("Client Name")
productName = st.sidebar.text_input("Product Name")

##########################################
#----------------------------------------#
#            Control Panel               #
#----------------------------------------#
##########################################


# Loading up the template document 
document = Document('UBR.docx')

## Setting up the bullet point (It may not be used)
for i in range(5):  # Adjust range for as many levels as you need
  try:
      style = document.styles.add_style(f'List Bullet {i}', document.styles['Normal'].type)
  except:
      style = document.styles[f'List Bullet {i}']
  style.paragraph_format.left_indent = Pt(18 * i)  # 36 points = 0.5 inches
  style.paragraph_format.first_line_indent = Pt(0)  # 18 points = 0.25 inches
  style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
  style.font.size = Pt(11)

## Setting up the width of columns in MBR steps
def set_col_widths(table):
    widths = (Inches(0.47), Inches(4.69), Inches(0.97), Inches(0.67), Inches(0.67))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

## Overall Font in the document
style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
#font.size = Pt(11)

## Spacing after and before table
def remove_table_spacing(doc):
    # Iterate through each paragraph in the document
    for paragraph in doc.paragraphs:
        # Check if the paragraph contains a table
        if paragraph._p.xml.startswith('<w:tbl'):
            # Access the paragraph's paragraph format
            paragraph_format = paragraph.paragraph_format
            # Set the spacing before and after the table to zero
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)


def set_vertical_cell_direction(cell: _Cell, direction: str):
    # direction: tbRl -- top to bottom, btLr -- bottom to top
    assert direction in ("tbRl", "btLr")
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    textDirection = OxmlElement('w:textDirection')
    textDirection.set(qn('w:val'), direction)  # btLr tbRl
    tcPr.append(textDirection)

def format_cell(cell, alignment, vertical_alignment, font_size):
    paragraph = cell.paragraphs[0]
    for run in paragraph.runs:
        run.font.size = Pt(font_size)  # Adjust the font size
        paragraph.alignment = alignment
    cell.vertical_alignment = vertical_alignment

##########################################
#----------------------------------------#
#      Step 1: General Info              #
#----------------------------------------#
##########################################


## Access the tables in word file 

## Iterate over each table in the Word Doc.

## Structure of the template
#Table 1: Section1 - Document Approval and Review
#Table 2: Section2 - General Information - Table of Content:
#Table 3: Section3 - Reference Information - Referenced Documents 
#Table 4: Section4 - Primary Packaging Equipment List 
#Table 5: Section5 - Primary Packaging Operations #<----- This is where we need to add input via Python


tablecounter1 = 0 #<---- This will count a number of tables being processed.
st.markdown("### Setting up Primary Packaging Operations")

with st.expander("Primary Packaging Operations"):
   col_ppo1, col_ppo2, col_ppo3 = st.columns(3)

   with col_ppo1:
      st.caption("Fill Count Input")
      fillcount = st.text_input("Fill count per bottle")
      fillcountref = st.text_input("Reference (ex. PSIS-Sec X) (ref ID 1)")
   with col_ppo2:
      st.caption("Total Bottle Required Input")
      totalbottle = st.text_input("Total Bottles Required")
      totalbottleref = st.text_input("Reference (ex. PSIS-Sec X) (ref ID 2)")
   with col_ppo3:
      st.caption("Enter the verification reference for Wipotec Scale")
      verWipotecref = st.text_input("Reference (ex. PSIS-Sec X) (ref ID 3)")

table = document.tables[4]
cells = [table.cell(13,2),table.cell(14,2),table.cell(16,3)]
texts = ['Fill Count per\n Bottle\n'+fillcount+'\n'+fillcountref,
         'Total Bottles\n Required\n'+totalbottle+'\n'+totalbottleref,
         verWipotecref]
for cell, text in zip(cells, texts):
         cell.text = text

# Format cells - Fontsize = 10
format_cell(cells[0], WD_PARAGRAPH_ALIGNMENT.CENTER, WD_ALIGN_VERTICAL.CENTER,10)
format_cell(cells[1], WD_PARAGRAPH_ALIGNMENT.CENTER, WD_ALIGN_VERTICAL.CENTER,10)
format_cell(cells[2], WD_PARAGRAPH_ALIGNMENT.CENTER, WD_ALIGN_VERTICAL.CENTER,10)


##########################################
#----------------------------------------#
#           Step 2: Packaging            #
#----------------------------------------#
##########################################

# This will be overall step number in the first column of MBR
StepNum = 0

# This will prevent error where if-statement generates error as undefined variable 
# if the step is not selected
TableFormat = sachet = canister = cotton = sealer = additional1 = False
cartoning = topsert = sidesert = shipper = bundling = additional2 = False

# Setting up the interface frame
st.markdown("### List of Processes")
col1, col2 = st.columns(2)

##################################
### Step 2-A Primary Pacakging ###
##################################

#################################################################################
### -------------- Primary Packaging -------------------

with col1:
   st.markdown('#### Primary Packaging')
   primary = st.checkbox("Primary Packaging")

   # Define all checkbox variables first
   
   if primary:
      ###--- Primary Packaging related list
      st.divider()
      sachet = st.checkbox("Sachet")
      canister = st.checkbox("Canister")
      cotton = st.checkbox("Cotton Filler")
      sealer = st.checkbox("Sealer")

      # break the page 
      document.add_page_break()

      subtitle = document.add_paragraph()
      run = subtitle.add_run('Part II: Primary Packaging')
      run.bold = True
      run.font.size = Pt(14)
      subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
      subtitle.paragraph_format.space_after = 0
      subtitle.paragraph_format.space_before = 0
      i = 0

st.divider()
st.markdown('### Process Control Panel')

##################################################################
if not primary:
   st.caption("Primary is not Selected")
   st.divider()


## Structure of the template
#Table 6: Primary Material
#Table 7: Packaging Material
#Table 8: Equipment List


if primary:
   st.markdown('#### Primary Packaging Step Selection')
   st.markdown('##### Primary Material Information')

   ############################
   ## Primary Material Setup ##
   ############################

   table = document.tables[5]

   numberProds = st.checkbox("More than one Item Number?")
   col_pri1, col_pri2, col_pri3 = st.columns(3)

   ### Item Number
   with col_pri1:
      
      if numberProds:
         prodItemNo1 = st.text_input("#1 Item No.")
         prodItemNo2 = st.text_input("#2 Item No.")
         ProdItemNo = prodItemNo1 + '\nor \n' + prodItemNo2
      else:
         ProdItemNo = st.text_input("Item No.")

   ### Name of product
   with col_pri2:
      productName = st.text_input("Name of Product")

   ### Theoretical Amount
   with col_pri3:
      theo_spec = st.text_input("Theoretical Amount required")
     
   ## Adding to pre-existing table
   
   cells = [table.cell(1,col) for col in range(0,3)]
   texts = [ProdItemNo,productName,theo_spec]
   for cell, text in zip(cells, texts):
         cell.text = text

   # Format cells - Fontsize = 12
   format_cell(cells[0], WD_PARAGRAPH_ALIGNMENT.CENTER, WD_ALIGN_VERTICAL.CENTER,12)
   format_cell(cells[1], WD_PARAGRAPH_ALIGNMENT.LEFT, WD_ALIGN_VERTICAL.CENTER,12)
   format_cell(cells[2], WD_PARAGRAPH_ALIGNMENT.CENTER, WD_ALIGN_VERTICAL.CENTER,12)

   #[Note]: I may need to make a section where we may have more than one primary material (not really)
   ############################################################

   ####################################
   ## Packaging Material Information ##
   ####################################

   table = document.tables[6]
   
   st.divider()
   st.markdown('##### Primary Packaging Material Information')
   noOfmaterials = st.number_input("Enter a number of packaging materials",min_value = 3, max_value = 10, value = 3)
   iter1 = int(noOfmaterials)
   
   with st.expander("Primary Packaging Material Info Entry"):
      col_pri4, col_pri5, col_pri6, col_pri7 = st.columns([1,1,2,1])

      itemNoInput1 = []
      itemNoInput2 = []
      matNameInput = []
      theoInput = []

      with col_pri4:
         for i in np.arange(iter1):
            itemNo1 = st.text_input(f'1st Item No. for Mat. No. {i+1}')
            itemNoInput1.append(itemNo1)

      with col_pri5:
         for i in np.arange(iter1):
            itemNo2 = st.text_input(f'2nd Item No. for Mat. No. {i+1} (if none, type N/A)')
            itemNoInput2.append(itemNo2)        

      with col_pri6:
         for i in np.arange(iter1):
            matName = st.text_input(f'Name for Mat. No. {i+1}')
            matNameInput.append(matName)


      with col_pri7:
         for i in np.arange(iter1):
            theoAmt = st.text_input(f'Theoretical Amount for Mat. No. {i+1}')
            theoInput.append(theoAmt)

      if iter1 > 3:
         for i in np.arange(iter1 - 3):
            row_cells = table.add_row().cells
            cell = table.cell(4+i,0)
            cell.text = 'Circle\nItem\n#(s)'
            paragraph = cell.paragraphs[0]
            run = paragraph.runs
            for run in paragraph.runs:
               paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_vertical_cell_direction(cell, 'btLr')


      for i,j in enumerate(itemNoInput1):
         if itemNoInput2[i] == 'N/A':
            itemNoInput = j
         else:
            itemNoInput = j+'\n and/or\n'+itemNoInput2[i]

         cells = [table.cell(i+1,col) for col in range(1,4)]
         texts = [itemNoInput, matNameInput[i],theoInput[i]]
         for cell, text in zip(cells, texts):
            cell.text = text

         # Format cells - Fontsize = 12
         format_cell(cells[0], WD_PARAGRAPH_ALIGNMENT.CENTER, WD_ALIGN_VERTICAL.CENTER,12)
         format_cell(cells[1], WD_PARAGRAPH_ALIGNMENT.LEFT, WD_ALIGN_VERTICAL.CENTER,12)
         format_cell(cells[2], WD_PARAGRAPH_ALIGNMENT.CENTER, WD_ALIGN_VERTICAL.CENTER,12)

         # Set row height
         table.rows[i+1].height = Inches(0.66)

   
   ###########################
   ## Equipment Information ##
   ###########################

   table = document.tables[7]
   
   st.divider()
   st.markdown('##### Primary Packaging Equipment Information')

   # Note: I should create a dictionary
   equiplist1 = {'Bottle Unscrambler':'ILS-1',
                 'Line Control':'Conveyor',
                 'Uniline':'IMA'}
   equiplist2 = {'Surekap Re-torquer':'SK600',
                 'Induction Sealer':'LM5412-T67',
                 'IMADA Torque Tester':'N/A'}
   equiplist3 = {'Wipotec Weight Checker':'N/A',
                 'Swiftcheck Tablet Capsule Counter':'N/A'}

   allequiplist = {**equiplist1, **equiplist2, **equiplist3}

   equipselected = []

   with st.expander("Primary Packaging Equipment List", expanded=True):
      col_ppe1, col_ppe2, col_ppe3 = st.columns(3)

      with col_ppe1:
         for i, equip in enumerate(list(equiplist1.keys())):
            option = st.checkbox(equip,value=False)
            if option:
               equipselected.append(equip)
       
      with col_ppe2:
         for i, equip in enumerate(list(equiplist2.keys())):
            option = st.checkbox(equip,value=False)
            if option:
               equipselected.append(equip)
      
      with col_ppe3:
         for i, equip in enumerate(list(equiplist3.keys())):
            option = st.checkbox(equip,value=False)
            if option:
               equipselected.append(equip)

   for i, equip in enumerate(equipselected, start=1):
      # Assign value to cells
      cells = [table.cell(i,col) for col in range(0,3)]
      texts = ['1', equip, allequiplist[equip]]
      
      for cell, text in zip(cells, texts):
         cell.text = text

      format_cell(cells[0], WD_PARAGRAPH_ALIGNMENT.CENTER, WD_ALIGN_VERTICAL.CENTER,12)
      format_cell(cells[1], WD_PARAGRAPH_ALIGNMENT.LEFT, WD_ALIGN_VERTICAL.CENTER,12)
      format_cell(cells[2], WD_PARAGRAPH_ALIGNMENT.CENTER, WD_ALIGN_VERTICAL.CENTER,12)

      # If not the last element, add a new row for the next iteration
      if i != len(equipselected):
         table.add_row()

   ###########################
   ##       Line Check      ##
   ###########################

   table = document.tables[9]
   cell = table.cell(6,1)
   paragraph = cell.add_paragraph()
   text = '\nRecord the batch number and quantity of '+productName+' available in the spaces provided.\n'
   run = paragraph.add_run(text)
   run.font.size = Pt(12)

   table = document.tables[10]
   for i in [1,3,5]:
      cell = table.cell(i,1)
      paragraph = cell.add_paragraph()
      text1 = '\nCollect one hundred (100) '+productName+' from the beginning of the bulk product allocated for this batch and printweigh (in grams) using the space provided. Record the scale number in the space provided.\n'
      text2 = '\nNote: All product used for the 100 ct. weights are to be returned to bulk product.\n'
      run1 = paragraph.add_run(text1)
      run2 = paragraph.add_run(text2)
      run2.bold = True
      run1.font.size = Pt(12)
      run2.font.size = Pt(12)
   

#################################################################################

## ---- Secondary Packaging ------------------------------

with col2:
   st.markdown('#### Secondary Packaging')
   secondary = st.checkbox("Secondary Packaging")
      
   if secondary:
      ###--- Secondary Packaging related list
      st.divider()
      cartoning = st.checkbox("Cartoning")
      sidesert = st.checkbox("Sidesert")
      topsert = st.checkbox("Topsert")
      bundling = st.checkbox("Bundling")
      shipper = st.checkbox("Shipper")

      # break the page 
      document.add_page_break()

   
      subtitle = document.add_paragraph()
      run = subtitle.add_run('Part III: Secondary Packaging')
      run.bold = True
      run.font.size = Pt(14)
      subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
      subtitle.paragraph_format.space_after = 0
      subtitle.paragraph_format.space_before = 0

##################################################################
if not secondary:
   st.caption("Secondary is not Selected")
   st.divider()

if secondary:
   st.markdown('#### Secondary Packaging Step Selection')





##########################################
#----------------------------------------#
#----------------------------------------#
##########################################
# Save the document
st.sidebar.header("**Step 7: Download Ready**")
if st.sidebar.checkbox("Check this box if the draft is ready"):
   with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
       document.save(tmp.name)
       tmp.seek(0)

       # Create a button to download the docx file
       st.sidebar.download_button(
           label="Download .docx file",
           data=tmp.read(),
           file_name="UBRDraft.docx",  ##<----- Make sure you change this to customizable command
           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
       )

   # Remove temporary file
   os.unlink(tmp.name)

st.sidebar.header('**RESET**')
if st.sidebar.checkbox("Ready to reset"):
   btn = st.sidebar.button("RESET")
   if btn:
      st.experimental_rerun()


